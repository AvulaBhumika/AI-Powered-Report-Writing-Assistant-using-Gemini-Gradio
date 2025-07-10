import os
import gradio as gr
import google.generativeai as genai
from datetime import datetime
from docx import Document

import os
os.environ["GEMINI_API_KEY"] = 'API_KEY_HERE'

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel("gemini-2.0-flash")

def save_to_docx(content: str, filename: str) -> str:
    doc = Document()
    doc.add_heading("Generated Report", level=1)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    return filename

def generate_report(insight_type: str, raw_data: str):
    """
    Generates report based on insight_type using Gemini model
    """
    prompt_map = {
        "dealer_summary": f"""
You are a market research assistant.

Based on the following dealer data, generate a region-wise dealer summary.
Highlight:
- Number of dealers
- Services provided (spares, service, both)
- Notable patterns or missing coverage

Raw Dealer Data:
{raw_data}
""",
        "customer_voice_snapshot": f"""
You are an AI assistant.

Summarize the following customer feedback:
- Extract overall sentiment (positive/neutral/negative)
- List repeated complaints or praises
- Identify customer suggestions

Raw Comments:
{raw_data}
""",
        "research_digest": f"""
You're helping generate a weekly/monthly automotive research digest.

Based on the following raw notes, generate a summary that includes:
- Key trends or patterns observed
- Brand/dealer-specific highlights
- Noteworthy changes or emerging concerns

Collected Data:
{raw_data}
"""
    }

    if insight_type not in prompt_map:
        return "Invalid type", None

    response = model.generate_content(prompt_map[insight_type])
    report_text = response.text
    filename = f"{insight_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = save_to_docx(report_text, filename)
    return report_text, filepath

# Gradio Interface
with gr.Blocks(title="Report Generation Assistant") as demo:
    with gr.Row():
        with gr.Column(scale=1):
            gr.Image(value="MS.png", width=100, show_label=False)
        with gr.Column(scale=5):
            gr.Markdown("### 🧾 **MintSurvey AI Report Generation Assistant**", elem_id="title-text")

    gr.Markdown("Use this tool to generate structured summaries from raw automotive data.")


    with gr.Row():
        insight_type = gr.Radio(
            choices=["dealer_summary", "customer_voice_snapshot", "research_digest"],
            label="Select Report Type",
            value="dealer_summary"
        )

    raw_data_input = gr.Textbox(
        label="Paste Raw Data Here",
        placeholder="e.g., Dealer: SK Automobiles, Bangalore – Offers: Service & Spares\n...",
        lines=10
    )

    generate_btn = gr.Button("🚀 Generate Report")
    report_output = gr.Textbox(label="Generated Report", lines=20)
    download_link = gr.File(label="📥 Download Report (.docx)")

    def run_generation(insight_type, raw_data):
        output_text, filepath = generate_report(insight_type, raw_data)
        return output_text, filepath

    generate_btn.click(
        fn=run_generation,
        inputs=[insight_type, raw_data_input],
        outputs=[report_output, download_link]
    )

if __name__ == "__main__":
    demo.launch()
